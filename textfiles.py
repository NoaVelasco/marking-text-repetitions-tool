#!/usr/bin/env python3.10
# Python 3.10.2 UTF-8
# Copyright (c) 2023, Noa Velasco

import re
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class Coinsidensias:
    """
    Relationed functions in order to compare words and change the similar ones.
    The output is a copy of the text, but with markdown format highlight.
    """

    def __init__(self, pathfile):
        self.path = Path(pathfile)
        self.name = pathfile.name
        self.simplename = pathfile.stem
        self.ext = pathfile.suffix
        # This gives a list with the words of the document:
        self.text = self.extract_text()
        self.new_text = []

    # -------------------------WORKING WITH FILES--------------------------------------

    def extract_text(self):
        """Determines what kind of file and choose the extract mode."""
        if self.ext == ".docx":
            return self.filedocx()
        return self.filetxt()

    def filedocx(self):
        """Extracts the contents of the docx file and returns a list of words."""
        print(f"«{self.name}» es un docx")

        doc_file = Document(self.path)
        doc_text = []

        for para in doc_file.paragraphs:
            if "Heading" in para.style.name:
                doc_text.append("# ")
            # Separates words and no-words/characters/spaces…
            prestring = re.split(r"([\W]+)", para.text)
            for word in prestring:
                doc_text.append(word)
            doc_text.append("  \n")

        return doc_text

    def filetxt(self):
        """Extracts the contents of the text file and returns a list of words."""
        print(f"«{self.name}» es un archivo de texto")

        with open(self.path, "r", encoding="utf-8") as f:
            txt_file = f.read()

        return re.split(r"([\W]+)", txt_file)

    # -------------------------WORKING WITH WORDS--------------------------------------

    def highlight(self, word):
        """Apply markdown format highlight adding a special character."""
        # "Ç" is an unused character, so I can search it safely later.
        # I know there is a way to do this with @decorators, but it's for my future me.
        return f"=={word}==Ç"

    def replacing(self, word):
        """Spanish words has special characters, most of them vowels with «tilde».
        This function replaces the word with a safe version only for comparing."""
        tildes = {
            "á": "a",
            "é": "e",
            "í": "i",
            "ó": "o",
            "ú": "u",
            "Á": "A",
            "E": "E",
            "Í": "I",
            "Ó": "O",
            "Ú": "U",
        }

        wildcard = []
        for l in word:
            if l in tildes:
                l = l.replace(l, tildes[l])
            wildcard.append(l.lower())
        word = "".join(wildcard)

        return word

    def scanwords(self, word1, word2):
        """Compares 2 words: if they have a matching nr of letters,
        return True, otherwise return False"""
        word1 = self.replacing(word1)
        word2 = self.replacing(word2)

        # minimum letters for word
        MIN = 5
        pattern = r"[a-zA-ZñÑ]{5}"
        x = re.search(pattern, word1)
        y = re.search(pattern, word2)
        if x is None:
            return False
        if y is None:
            return False

        for n in range(len(word1) - MIN + 1):
            searching = word1[0 + n : MIN + n]
            searching = re.search(searching, word2)
            if searching is not None:
                return True
        return False

    # -------------------------DOING THE MAGIC--------------------------------------

    def activate(self):
        """Scan a text word by word, then compare they with a number of words before
        and after them. If similar, append a highlighted version in the new text list"""

        # To give a list of exceptions; these words will not be taken into account:
        # except_file = Path(r'path/exceptions.txt')
        # with open(except_file, encoding='utf-8') as ex:
        #     exceptions = ex.read().splitlines()

        for i, word1 in enumerate(self.text):
            reqs = False
            words_cnt = 0
            word2 = ""
            highld = False
            is_word = r"([\W]+)"

            while reqs is False:
                # Spaces are included, so -60/+121 for nearby 20 words around.
                y = i - 60 + words_cnt
                scape_word = re.search(is_word, word1)

                # If an exception list was given before, activate this.
                # if scape_word is not None or word1 in exceptions:
                if scape_word is not None:
                    # meets requirements to exit the loop: regular append.
                    reqs = True
                elif y == i:
                    # same word instance, so next word.
                    words_cnt += 1
                elif y < 0:
                    # avoid index errors at the beginning.
                    words_cnt += abs(y)
                elif words_cnt == 121 or y >= len(self.text):
                    # end of this comparison.
                    reqs = True
                else:
                    # meets requirements to scanwords.
                    word2 = self.text[y]
                    scape_word = re.search(is_word, word2)

                    # If this is not a real word, go next.
                    if scape_word is not None:
                        words_cnt += 1
                    else:
                        compare = self.scanwords(word1, word2)

                        if compare is True:
                            self.new_text.append(self.highlight(word1))
                            # positive match: highlight this word, append,
                            # but avoid a new append out of the loop (with this condition True).
                            highld = True
                            reqs = True
                        else:
                            words_cnt += 1
            if highld is False:
                self.new_text.append(word1)

    # -------------------------WRITTING THE VERSIONS--------------------------------------

    def write_docx(self, docx):
        '''Pure python-docx magic and sorcery; and a bit of anger and frustration.'''
        pre_py_docx = []

        # It can be easier to detect a highlighted word if is a list.
        for word in self.new_text:
            if "Ç" in word:
                word = word.split("==")
            pre_py_docx.append(word)

        style = docx.styles
        par_style = style.add_style("custom", WD_STYLE_TYPE.PARAGRAPH)

        par_style.paragraph_format.first_line_indent = Inches(0.25)
        par_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par_style.paragraph_format.space_before = Pt(0)
        par_style.paragraph_format.space_after = Pt(0)
        par_font = par_style.font
        # I recommend Bookerly font because it's very readable. But NTR is universal.
        par_font.name = "Times New Roman"
        par_font.size = Pt(12)

        h1_style = style.add_style("h1", WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = style["Heading 1"]
        font = h1_style.font
        font.name = "Times New Roman"
        font.size = Pt(14)

        # @nota IDK why "highlight_color" has a weird behaviour.
        # I can't set this in it's own style and it can disappear without apparent reason.
        char_charstyle = style.add_style("highlight", WD_STYLE_TYPE.CHARACTER)
        char_font = char_charstyle.font
        char_font.name = "Times New Roman"
        char_font.size = Pt(12)

        docx.add_heading("Versión docx con palabras resaltadas")

        sigue = docx.add_paragraph("", style="custom")
        run = sigue.add_run()

        for word in pre_py_docx:
            # This is the best solution with line breaks after a lot of testing.
            if word == "  \n":
                sigue = docx.add_paragraph("", style="custom")
            # it's a list, so highlight word[1].
            elif isinstance(word, list):
                run = sigue.add_run(word[1])
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            # it's a heading, so…
            elif word == "# ":
                sigue = docx.add_paragraph(style="h1")
                run = sigue.add_run()
            # regular text.
            else:
                run = sigue.add_run(word)
                run.font.name = "Times New Roman"

    def savedocx(self):
        '''Save a copy in docx format.'''
        doc_copy = Document()
        self.write_docx(doc_copy)

        copy_name = f"{self.path.parent}/{self.simplename}-py.docx"
        doc_copy.save(copy_name)
        print("El archivo ha sido copiado con cambios.")

    def savemd(self):
        '''Save a copy in markdown format.'''
        copy_name = f"{self.path.parent}/{self.simplename}-py.md"
        md_text = []
        for word in self.new_text:
            if "Ç" in word:
                word = word.replace("==Ç", "==")
            md_text.append(word)

        md_text = "".join(md_text)
        with open(copy_name, "w", encoding="utf-8") as txt:
            txt.write(md_text)
        print("El archivo ha sido copiado con cambios.")

    def writefiles(self):
        '''Once it has the new text, activates the functions and saves.'''
        self.activate()
        self.savedocx()
        self.savemd()
